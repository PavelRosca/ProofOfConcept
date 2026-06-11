from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/home/user/ProofOfConcept/rapoarte_1_4"

def set_font(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_period(doc, period_label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run(f"{period_label}: ")
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_body = p.add_run(text)
    run_body.bold = False
    run_body.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_title_block(doc, role_title, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{role_title} – {name}")
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    return p

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

def create_doc(filename, role_title, name, p1, p2, p3):
    doc = Document()
    set_margins(doc)

    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    add_title_block(doc, role_title, name)
    add_period(doc, "P1", p1)
    add_period(doc, "P2", p2)
    add_period(doc, "P3", p3)

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"  Salvat: {filename}")

# ─────────────────────────────────────────────────────────────────────────────
# 1. CERCETĂTOR — Alex Bere
# ─────────────────────────────────────────────────────────────────────────────
create_doc(
    filename="Cercetator_1_4_Alex_Bere.docx",
    role_title="Cercetător în informatică / Supervizor / Head of R&D",
    name="Alex Bere",
    p1=(
        "Definirea cadrului metodologic de evaluare a performanței sistemului integrat în cadrul "
        "subactivității 1.4. Elaborarea matricei indicatorilor de performanță per subsistem — robot "
        "terestru, platformă aeriană, platformă software de management și componentă de inteligență "
        "artificială —, cu stabilirea valorilor de referință corespunzătoare obiectivelor de inspecție "
        "autonomă ale proiectului. Inițierea cercetării soluțiilor de încărcare bazate pe energie "
        "solară, cu definirea parametrilor de analiză: necesarul energetic al echipamentelor mobile, "
        "variabilele meteorologice relevante pentru operarea în teren deschis și criteriile de "
        "dimensionare a infrastructurii de alimentare. Supervizarea fazei de start a subactivității și "
        "coordonarea distribuirii sarcinilor de evaluare între membrii echipei."
    ),
    p2=(
        "Supervizarea evaluării componentei de inteligență artificială pentru detecția panourilor "
        "fotovoltaice și a anomaliilor termice: validarea seturilor de date utilizate la antrenare, a "
        "procedurii de antrenare pe CPU și a metricilor de acuratețe obținute (mAP50, precizie, recall). "
        "Coordonarea analizei energetice a platformelor mobile în condiții meteorologice variate, cu "
        "accent pe impactul temperaturilor scăzute asupra capacității efective a acumulatorilor și "
        "implicațiile acestuia asupra planificării misiunilor de inspecție. Supervizarea integrării "
        "rezultatelor de performanță obținute individual de fiecare subsistem într-un model unitar de "
        "evaluare transversală, asigurând coerența metodologică între evaluările cantitative și cele "
        "conceptuale."
    ),
    p3=(
        "Coordonarea sintezei transversale a performanței sistemului integrat și formularea ajustărilor "
        "recomandate la nivelul algoritmilor de inteligență artificială și al specificațiilor de "
        "echipamente, pe baza observațiilor colectate în perioadele anterioare. Validarea analizei "
        "bilanțului energetic și a scenariilor de încărcare solară documentate în raport, inclusiv a "
        "dimensionărilor propuse pentru stațiile solare dedicate. Validarea raportului tehnic al "
        "subactivității 1.4 în ansamblu — structura, corectitudinea tehnică, coerența concluziilor și "
        "alinierea cu obiectivele proiectului —, precum și a direcțiilor de continuare a cercetării "
        "formulate pentru etapele următoare."
    )
)

# ─────────────────────────────────────────────────────────────────────────────
# 2. EXPERT BACK-END — Gangoș George
# ─────────────────────────────────────────────────────────────────────────────
create_doc(
    filename="Expert_BackEnd_1_4_Gangos_George.docx",
    role_title="Expert în Back-End Development",
    name="Gangoș George",
    p1=(
        "Definirea indicatorilor de performanță pentru platforma software de management în contextul "
        "subactivității 1.4: timpi de răspuns ai endpoint-urilor REST, debit de interogări per "
        "categorie de complexitate și comportamentul fondului de conexiuni la baza de date PostgreSQL. "
        "Analiza structurală a interogărilor cu risc de degradare a performanței — problema N+1 la "
        "lista sarcinilor și a tehnicienilor, interogările agregate ale tabloului de bord și "
        "interogarea de verificare a utilizatorului în middleware-ul de autentificare —, pornind de "
        "la arhitectura validată în subactivitatea 1.3. Documentarea bazei de referință pentru "
        "evaluare: stiva Fastify + Prisma + PostgreSQL, structura modulară pe domenii de business și "
        "modelul de autentificare JWT cu control al accesului bazat pe roluri."
    ),
    p2=(
        "Evaluarea detaliată a performanței platformei software în condiții de sarcină realistă: "
        "analiza timpilor de răspuns per categorie de endpoint — interogări simple față de interogări "
        "cu asocieri multiple —, estimarea debitului orientativ și identificarea punctelor de "
        "degradare în scenarii de utilizare concurentă. Documentarea absențelor identificate ca "
        "principale pârghii de optimizare: absența stratului de cache și absența compresiei HTTP, cu "
        "evaluarea impactului estimat al introducerii lor asupra capacității platformei. Analiza "
        "metricilor componentei de inteligență artificială pentru detecția defectelor integrate în "
        "fluxul backend: latența inferenței pe server, debitul de procesare a imaginilor și "
        "comportamentul arhitecturii de servire pe mai multe niveluri sub sarcină."
    ),
    p3=(
        "Documentarea recomandărilor de optimizare a platformei software pentru etapele următoare: "
        "implementarea unui strat de cache pentru interogările frecvente, activarea compresiei HTTP, "
        "calibrarea dimensiunii fondului de conexiuni și configurarea unui framework de benchmarking "
        "pentru validarea îmbunătățirilor. Contribuție la ajustările specificațiilor de echipamente "
        "privind infrastructura server necesară execuției accelerate a modelelor AI: evaluarea "
        "comparativă CPU multi-core față de GPU dedicat, cu referire la diferența de performanță "
        "documentată (8–15 FPS pe CPU față de 25–30 FPS cu accelerare TensorRT pe GPU). Validarea "
        "secțiunilor din raportul 1.4 referitoare la performanța platformei software și a componentei "
        "AI, asigurând consistența cu arhitectura documentată în subactivitatea 1.3."
    )
)

# ─────────────────────────────────────────────────────────────────────────────
# 3. EXPERT DRONE & AR — Pavel Rosca
# ─────────────────────────────────────────────────────────────────────────────
create_doc(
    filename="Expert_Drone_AR_1_4_Pavel_Rosca.docx",
    role_title="Expert în drone și realitate augmentată",
    name="Pavel Rosca",
    p1=(
        "Cercetarea soluțiilor de încărcare autonomă bazate pe energie solară pentru platforma aeriană, "
        "pornind de la arhitectura DJI Cloud API integrată în subactivitatea 1.3. Analiza specificațiilor "
        "tehnice ale stațiilor de andocare comerciale compatibile cu drona DJI Matrice — inclusiv "
        "mecanismele de aterizare autonomă, condițiile de operare și cerințele de alimentare ale "
        "stațiilor —, precum și evaluarea modului în care acestea se integrează cu platforma cloud "
        "privată on-premise implementată anterior. Analiza necesarului energetic al dronei: capacitatea "
        "acumulatorului, autonomia de zbor nominală și realistă, timpii de reîncărcare în mod rapid și "
        "impactul temperaturii asupra capacității efective a bateriei în sezonul rece."
    ),
    p2=(
        "Evaluarea performanței subsistemului aerian în condiții meteorologice variate: analiza "
        "impactului vântului asupra consumului energetic și a stabilității zborului, cuantificarea "
        "reducerii autonomiei la temperaturi scăzute, identificarea pragurilor operaționale în condiții "
        "de precipitații sau rafale. Evaluarea performanței fluxurilor de comunicare ale platformei "
        "aeriene stabilite în subactivitatea 1.3: frecvența telemetriei MQTT, latența comenzilor de "
        "zbor și latența video end-to-end prin protocolul WebRTC în scenarii operaționale reprezentative. "
        "Documentarea scenariilor de încărcare solară aplicabile platformei aeriene: stație centrală "
        "unică față de stații distribuite pe suprafața parcului, cu evaluarea numărului de cicluri de "
        "reîncărcare necesare pe zi în funcție de intensitatea misiunilor."
    ),
    p3=(
        "Documentarea rezultatelor cercetării privind autonomia energetică a dronei și a soluțiilor de "
        "andocare solară, cu concluzii privind fezabilitatea operării continue pe un parc fotovoltaic de "
        "1 MW. Contribuție la ajustările specificațiilor de echipamente recomandate în raportul 1.4: "
        "evaluarea necesității achiziției unei stații de andocare dedicate, analiza numărului optim de "
        "puncte de andocare în funcție de suprafața parcului și de autonomia de zbor disponibilă, "
        "precum și evaluarea funcționalității de control direct la distanță (DRC) în raport cu "
        "scenariile de misiune autonomă. Validarea bilanțului energetic al platformei aeriene și a "
        "estimărilor privind ponderea consumului flotei în producția totală a parcului fotovoltaic."
    )
)

# ─────────────────────────────────────────────────────────────────────────────
# 4. EXPERT DB & BIG DATA — Alexandru Indries
# ─────────────────────────────────────────────────────────────────────────────
create_doc(
    filename="Expert_Db_Big_Data_1_4_Alexandru_Indries.docx",
    role_title="Expert în big data și baze de date",
    name="Alexandru Indries",
    p1=(
        "Proiectarea cadrului de colectare a datelor de performanță pentru sistemul integrat în cadrul "
        "subactivității 1.4: definirea structurii de stocare a metricilor de telemetrie, a indicatorilor "
        "de performanță ai modelelor de inteligență artificială și a datelor de consum energetic al "
        "platformelor mobile. Analiza volumelor de date generate de subsisteme în operare continuă — "
        "date LiDAR de la robotul terestru, telemetrie MQTT și înregistrări video de la platforma "
        "aeriană, metrici API de la platforma software — și evaluarea cerințelor de stocare și "
        "procesare pentru un scenariu de operare pe un parc fotovoltaic de 1 MW. Documentarea "
        "fluxurilor de date relevante pentru evaluarea performanței, pornind de la arhitectura "
        "validată în subactivitatea 1.3."
    ),
    p2=(
        "Implementarea mecanismelor de colectare și agregare a datelor de performanță: gestionarea "
        "metricilor de telemetrie cu limitare temporală pentru scrierile în baza de date PostgreSQL, "
        "analiza comportamentului cache-ului Redis pentru stările curente ale dispozitivelor în "
        "condiții de sarcină ridicată și evaluarea consistenței datelor de performanță colectate din "
        "surse multiple. Analiza performanței modelelor de date sub sarcini realiste: evaluarea "
        "eficienței indexurilor existente pe coloanele frecvent interogate, identificarea interogărilor "
        "cu cost ridicat de execuție și documentarea comportamentului fondului de conexiuni în "
        "scenarii de concurență. Validarea integrității și trasabilității datelor de telemetrie "
        "colectate de la platforma aeriană prin DJI Cloud API."
    ),
    p3=(
        "Generarea analizei de performanță pe baza datelor colectate: procesarea metricilor de "
        "telemetrie, validarea consistenței datelor de consum energetic ale platformelor mobile și "
        "documentarea indicatorilor de performanță confirmați prin măsurători instrumentate. Contribuție "
        "la analiza bilanțului energetic pentru scenariile de încărcare solară prin structurarea și "
        "procesarea datelor de referință privind producția parcurilor fotovoltaice și necesarul "
        "energetic al echipamentelor — robot terestru și dronă —, utilizate în dimensionarea "
        "stațiilor solare din raportul 1.4. Documentarea recomandărilor privind arhitectura de date "
        "și a direcțiilor de optimizare a accesului la informații pentru etapele următoare ale "
        "proiectului."
    )
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. EXPERT SECURITATE CIBERNETICĂ — Vlad Botis
# ─────────────────────────────────────────────────────────────────────────────
create_doc(
    filename="Expert_Securitate_Cibernetica_1_4_Vlad_Botis.docx",
    role_title="Expert în securitate cibernetică",
    name="Vlad Botis",
    p1=(
        "Evaluarea overhead-ului mecanismelor de securitate implementate în subactivitatea 1.3 asupra "
        "performanței generale a sistemului integrat. Analiza costului computațional al validării "
        "token-urilor JWT per cerere autentificată, al funcțiilor de derivare a parolelor — argon2id "
        "pe platforma software, bcrypt pe platforma aeriană — și al interogărilor de verificare a "
        "stării contului executate în middleware-ul de autentificare. Documentarea relației dintre "
        "măsurile de securitate deja implementate și bugetul de performanță disponibil per subsistem, "
        "cu scopul de a integra dimensiunea de securitate în modelul de evaluare a performanței din "
        "subactivitatea 1.4."
    ),
    p2=(
        "Benchmarking-ul overhead-ului de securitate pe componentele platformei: evaluarea impactului "
        "middleware-ului de autentificare asupra timpilor de răspuns API, documentarea costului de "
        "procesare al validării datelor de intrare prin pipeline-ul global și al mecanismelor de "
        "logging structurat active. Analiza implicațiilor de performanță ale activării TLS pe canalele "
        "de comunicare identificate ca necriptate în auditul din subactivitatea 1.3 — HTTP, MQTT, "
        "WebSocket —: estimarea overhead-ului adăugat de handshake-ul TLS și de criptarea continuă a "
        "traficului operațional, cu referire la impactul estimat asupra latenței comenzilor de zbor și "
        "a frecvenței telemetriei. Evaluarea overhead-ului potențial al activării DDS Security pe "
        "robotul terestru și impactul său asupra frecvenței de procesare a norului de puncte LiDAR."
    ),
    p3=(
        "Documentarea compromisurilor acceptabile dintre securitate și performanță pentru fiecare "
        "subsistem în parte: definirea nivelurilor de overhead tolerabile fără impact semnificativ "
        "asupra misiunilor de inspecție autonomă, cu referire la constrângerile de timp real ale "
        "robotului și la latența critică a comenzilor de zbor pentru dronă. Contribuție la ajustările "
        "specificațiilor de echipamente din perspectiva securității: evaluarea cerințelor de procesare "
        "pentru activarea criptării pe canalele operaționale și a necesității unui modul hardware "
        "dedicat funcțiilor criptografice pe platforma edge. Validarea secțiunilor din raportul 1.4 "
        "privind relația performanță–securitate și a recomandărilor de ajustare formulate în "
        "concluzii, asigurând coerența cu vulnerabilitățile documentate și cu măsurile de remediere "
        "propuse în subactivitatea anterioară."
    )
)

print("\nToate fișierele au fost generate cu succes.")
